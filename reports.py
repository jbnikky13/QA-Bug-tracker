"""
reports.py — generates QA/accessibility bug reports (PDF and DOCX) formatted
to match the structure expected by bug bounty platforms (HackerOne/Bugcrowd
style): Title, Summary, Severity, Affected Asset, Steps to Reproduce,
Proof of Concept, Impact, Suggested Remediation, References.

Expected bug dict fields (see scanner.py's _bug()):
    id, title, severity, priority, type, page, message, evidence,
    wcag, selector, html_snippet, remediation, steps, expected, actual,
    help_url, screenshot
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image,
    PageBreak, HRFlowable
)

SEVERITY_COLOR = {
    "Critical": "C00000",
    "High": "E36C09",
    "Medium": "BF8F00",
    "Low": "548235",
}
SEVERITY_COLOR_RGB = {sev: colors.HexColor(f"#{hexval}") for sev, hexval in SEVERITY_COLOR.items()}

# Rough CVSS-style qualitative band, shown alongside severity the way most
# bounty platforms pair a label with a numeric-ish range for triage.
SEVERITY_CVSS_BAND = {
    "Critical": "9.0 – 10.0",
    "High": "7.0 – 8.9",
    "Medium": "4.0 – 6.9",
    "Low": "0.1 – 3.9",
}


def _summary_counts(bugs):
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    for b in bugs:
        counts[b.get("severity", "Low")] = counts.get(b.get("severity", "Low"), 0) + 1
    return counts


def _bug_title(b):
    return b.get("title") or f'{b.get("type","Finding")} on {b.get("page","")}'


# ---------------------------------------------------------------- DOCX -----

def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def make_docx(result, out_path, researcher=None, program=None):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    bugs = result["bugs"]
    pages = result["pages"]
    counts = _summary_counts(bugs)
    generated = datetime.now().strftime("%B %d, %Y %H:%M")
    target = result.get("target", "")

    # --- Cover / submission header ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Vulnerability & QA Findings Report")
    run.font.size = Pt(26)
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(target)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    header_table = doc.add_table(rows=0, cols=2)
    header_table.style = "Light List"
    for label, value in [
        ("Program / Target", program or target),
        ("Submitted by", researcher or "Automated QA Scan"),
        ("Date", generated),
        ("Assets in Scope", target),
        ("Pages Tested", str(len(pages))),
        ("Total Findings", str(len(bugs))),
    ]:
        row = header_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = value

    doc.add_page_break()

    # --- Summary of findings (submission-list style) ---
    doc.add_heading("Summary of Findings", level=1)
    sum_table = doc.add_table(rows=1, cols=4)
    sum_table.style = "Light Grid Accent 1"
    hdr = sum_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "ID", "Title", "Severity", "Status"
    for b in bugs:
        row = sum_table.add_row().cells
        row[0].text = b.get("id", "")
        row[1].text = _bug_title(b)
        row[2].text = b.get("severity", "")
        row[3].text = "Open"
        _set_cell_shading(row[2], SEVERITY_COLOR.get(b.get("severity", "Low"), "808080"))
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        row[2].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    sev_table = doc.add_table(rows=1, cols=2)
    sev_table.style = "Light Grid Accent 1"
    hdr = sev_table.rows[0].cells
    hdr[0].text, hdr[1].text = "Severity", "Count"
    for sev in ["Critical", "High", "Medium", "Low"]:
        row = sev_table.add_row().cells
        row[0].text = sev
        row[1].text = str(counts.get(sev, 0))
        _set_cell_shading(row[0], SEVERITY_COLOR[sev])
        row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        row[0].paragraphs[0].runs[0].font.bold = True

    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "Findings were identified via an automated scan using a headless Chromium browser (Playwright). "
        "Checks covered broken links/images, HTTP errors, JavaScript console errors, failed network "
        "requests, layout/rendering defects, and accessibility violations assessed against WCAG 2.x "
        "success criteria via axe-core."
    )

    doc.add_page_break()

    # --- Individual findings, one per "submission" ---
    doc.add_heading("Findings", level=1)

    if not bugs:
        doc.add_paragraph("No issues were detected during this scan.")

    for b in bugs:
        h = doc.add_heading(f'{b.get("id","")}: {_bug_title(b)}', level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(SEVERITY_COLOR.get(b.get("severity", "Low"), "000000"))

        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Light Grid"

        def add_row(label, value):
            row = meta.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].font.bold = True
            row[1].text = str(value) if value else "N/A"

        add_row("Severity", f'{b.get("severity")} (CVSS-equivalent range: {SEVERITY_CVSS_BAND.get(b.get("severity",""), "N/A")})')
        add_row("Priority", b.get("priority"))
        add_row("Weakness / Category", b.get("type"))
        add_row("Affected Asset (URL)", b.get("page"))
        add_row("Affected Element", b.get("selector"))
        if b.get("wcag") and b.get("wcag") != "N/A":
            add_row("WCAG Reference", b.get("wcag"))

        doc.add_paragraph()
        p = doc.add_paragraph(); p.add_run("Summary").bold = True
        doc.add_paragraph(b.get("message", ""))

        p = doc.add_paragraph(); p.add_run("Steps to Reproduce").bold = True
        for line in (b.get("steps") or "").split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Number")

        p = doc.add_paragraph(); p.add_run("Proof of Concept").bold = True
        if b.get("html_snippet"):
            doc.add_paragraph(f'Element: {b["html_snippet"]}')
        doc.add_paragraph(f'Expected: {b.get("expected","")}')
        doc.add_paragraph(f'Actual: {b.get("actual","")}')
        if b.get("screenshot") and Path(b["screenshot"]).exists():
            try:
                doc.add_picture(b["screenshot"], width=Inches(5.5))
            except Exception:
                pass

        p = doc.add_paragraph(); p.add_run("Impact").bold = True
        doc.add_paragraph(
            f'This issue affects the {b.get("severity","").lower()}-priority quality of the affected page. '
            f'{b.get("actual","")} This may degrade user experience, functionality, or accessibility compliance '
            f'depending on deployment context.'
        )

        p = doc.add_paragraph(); p.add_run("Suggested Remediation").bold = True
        doc.add_paragraph(b.get("remediation", ""))

        if b.get("help_url"):
            p = doc.add_paragraph(); p.add_run("References").bold = True
            doc.add_paragraph(b.get("help_url"))

        doc.add_paragraph("_" * 90)

    doc.save(str(out_path))
    return str(out_path)


# ----------------------------------------------------------------- PDF -----

def make_pdf(result, out_path, researcher=None, program=None):
    bugs = result["bugs"]
    pages = result["pages"]
    counts = _summary_counts(bugs)
    generated = datetime.now().strftime("%B %d, %Y %H:%M")
    target = result.get("target", "")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleBig", parent=styles["Normal"], fontName="Helvetica-Bold",
                                  fontSize=24, leading=28, spaceAfter=6, alignment=1)
    sub_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=13, textColor=colors.HexColor("#444444"),
                                alignment=1, spaceAfter=10)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], spaceBefore=14, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=12, spaceAfter=4)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], spaceBefore=8, spaceAfter=3)
    body = styles["BodyText"]
    label = ParagraphStyle("Label", parent=body, fontName="Helvetica-Bold")

    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER,
                             topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                             leftMargin=0.8 * inch, rightMargin=0.8 * inch)
    story = []

    # Cover / submission header
    story.append(Spacer(1, 1.3 * inch))
    story.append(Paragraph("Vulnerability &amp; QA Findings Report", title_style))
    story.append(Paragraph(target, sub_style))

    header_data = [
        ["Program / Target", program or target],
        ["Submitted by", researcher or "Automated QA Scan"],
        ["Date", generated],
        ["Assets in Scope", target],
        ["Pages Tested", str(len(pages))],
        ["Total Findings", str(len(bugs))],
    ]
    header_table = Table(header_data, colWidths=[2.0 * inch, 4.3 * inch])
    header_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.lightgrey),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(header_table)
    story.append(PageBreak())

    # Summary of findings (submission-list style)
    story.append(Paragraph("Summary of Findings", h1))
    sf_data = [["ID", "Title", "Severity", "Status"]] + [
        [b.get("id", ""), Paragraph(_bug_title(b), body), b.get("severity", ""), "Open"] for b in bugs
    ]
    sf_table = Table(sf_data, colWidths=[0.7 * inch, 3.5 * inch, 0.9 * inch, 0.9 * inch], repeatRows=1)
    sf_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F2F2F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    for i, b in enumerate(bugs, start=1):
        sf_style.append(("BACKGROUND", (2, i), (2, i), SEVERITY_COLOR_RGB.get(b.get("severity", "Low"))))
        sf_style.append(("TEXTCOLOR", (2, i), (2, i), colors.white))
    sf_table.setStyle(TableStyle(sf_style))
    story.append(sf_table)
    story.append(Spacer(1, 0.2 * inch))

    sev_data = [["Severity", "Count"]] + [[s, str(counts.get(s, 0))] for s in ["Critical", "High", "Medium", "Low"]]
    sev_table = Table(sev_data, colWidths=[2.5 * inch, 2.5 * inch])
    sev_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F2F2F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
    ]
    for i, sev in enumerate(["Critical", "High", "Medium", "Low"], start=1):
        sev_style.append(("BACKGROUND", (0, i), (0, i), SEVERITY_COLOR_RGB[sev]))
        sev_style.append(("TEXTCOLOR", (0, i), (0, i), colors.white))
    sev_table.setStyle(TableStyle(sev_style))
    story.append(sev_table)

    story.append(Paragraph("Methodology", h2))
    story.append(Paragraph(
        "Findings were identified via an automated scan using a headless Chromium browser (Playwright). "
        "Checks covered broken links/images, HTTP errors, JavaScript console errors, failed network "
        "requests, layout/rendering defects, and accessibility violations assessed against WCAG 2.x "
        "success criteria via axe-core.", body))
    story.append(PageBreak())

    # Individual findings
    story.append(Paragraph("Findings", h1))
    if not bugs:
        story.append(Paragraph("No issues were detected during this scan.", body))

    for b in bugs:
        sev = b.get("severity", "Low")
        heading_style = ParagraphStyle(f"H2_{sev}_{b.get('id','')}", parent=h2,
                                        textColor=SEVERITY_COLOR_RGB.get(sev, colors.black))
        story.append(Paragraph(f'{b.get("id","")}: {_bug_title(b)}', heading_style))

        meta_data = [
            ["Severity", f'{sev} (CVSS-equivalent: {SEVERITY_CVSS_BAND.get(sev,"N/A")})'],
            ["Priority", b.get("priority", "")],
            ["Weakness / Category", b.get("type", "")],
            ["Affected Asset (URL)", b.get("page", "")],
            ["Affected Element", b.get("selector", "N/A")],
        ]
        if b.get("wcag") and b.get("wcag") != "N/A":
            meta_data.append(["WCAG Reference", b.get("wcag")])
        meta_table = Table(meta_data, colWidths=[1.7 * inch, 4.6 * inch])
        meta_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.lightgrey),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.08 * inch))

        story.append(Paragraph("Summary", h3))
        story.append(Paragraph(b.get("message", ""), body))

        story.append(Paragraph("Steps to Reproduce", h3))
        for line in (b.get("steps") or "").split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), body))

        story.append(Paragraph("Proof of Concept", h3))
        if b.get("html_snippet"):
            story.append(Paragraph(f'<b>Element:</b> {b["html_snippet"]}', body))
        story.append(Paragraph(f'<b>Expected:</b> {b.get("expected","")}', body))
        story.append(Paragraph(f'<b>Actual:</b> {b.get("actual","")}', body))
        if b.get("screenshot") and Path(b["screenshot"]).exists():
            try:
                story.append(Spacer(1, 0.05 * inch))
                story.append(Image(b["screenshot"], width=5.2 * inch, height=3.0 * inch, kind="proportional"))
            except Exception:
                pass

        story.append(Paragraph("Impact", h3))
        story.append(Paragraph(
            f'This issue affects the {sev.lower()}-priority quality of the affected page. '
            f'{b.get("actual","")} This may degrade user experience, functionality, or accessibility '
            f'compliance depending on deployment context.', body))

        story.append(Paragraph("Suggested Remediation", h3))
        story.append(Paragraph(b.get("remediation", ""), body))

        if b.get("help_url"):
            story.append(Paragraph("References", h3))
            story.append(Paragraph(b.get("help_url"), body))

        story.append(Spacer(1, 0.12 * inch))
        story.append(HRFlowable(width="100%", color=colors.lightgrey))
        story.append(Spacer(1, 0.12 * inch))

    doc.build(story)
    return str(out_path)
